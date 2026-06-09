"""Generate a Word ads-performance report for a month. Default: June 2026.
Saves <project root>/Meta_Ads_Performance_<Month>_<Year>.docx."""
import io, re, sys
from pathlib import Path
from datetime import date, datetime

import duckdb
import pandas as pd
import requests
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SINCE, UNTIL = "2026-06-01", "2026-06-30"
MONTH_LABEL = "June 2026"
BLUE, GREEN, PURPLE, AMBER = "#3b82f6", "#10b981", "#8b5cf6", "#d08700"


def usd_to_aud():
    try:
        r = requests.get("https://api.frankfurter.app/latest",
                         params={"from": "USD", "to": "AUD"}, timeout=10)
        if r.ok:
            return float(r.json()["rates"]["AUD"])
    except Exception:
        pass
    return 1.52


FX = usd_to_aud()
con = duckdb.connect(str(ROOT / "data" / "migration_dashboard.duckdb"), read_only=True)

# ---- Meta spend/impr/clicks/leads (daily grain) ----
daily = con.execute(f"""
    SELECT date, account_label, campaign_name, SUM(spend) spend, SUM(impressions) impr,
           SUM(clicks) clicks, SUM(total_leads) leads
    FROM fact_meta_daily WHERE date BETWEEN DATE '{SINCE}' AND DATE '{UNTIL}'
    GROUP BY 1,2,3
""").fetchdf()
daily["spend_aud"] = daily["spend"] * FX
data_days = sorted(daily["date"].unique())
last_day = max(data_days).strftime("%b %d") if len(data_days) else "-"

# ---- GHL Meta-attributed leads + bookings + shows (unified-leads cohort) ----
sqltxt = ((ROOT / "dashboards" / "sql" / "executive_cards.sql").read_text(encoding="utf-8")
          + "\n" + (ROOT / "dashboards" / "sql" / "tab_cards.sql").read_text(encoding="utf-8"))
m = re.search(r'CREATE\s+OR\s+REPLACE\s+VIEW\s+vw_exec_unified_leads\s+AS\s+(.*?)(?=CREATE\s+OR\s+REPLACE\s+VIEW|\Z)',
              sqltxt, re.DOTALL | re.IGNORECASE)
ul = con.execute(m.group(1).strip().rstrip(";"), {"since": SINCE, "until": UNTIL}).fetchdf()
meta_ul = ul[ul["source_bucket"] == "Meta Paid"]
ghl_meta_leads = len(meta_ul)
ghl_booked = int(meta_ul["has_booking"].sum())
ghl_showed = int(meta_ul["showed"].sum())
con.close()

# ---- aggregates ----
spend = float(daily["spend_aud"].sum())
impr = float(daily["impr"].sum())
clicks = float(daily["clicks"].sum())
meta_leads = int(daily["leads"].sum())
ctr = (clicks / impr) if impr else 0
cpc = (spend / clicks) if clicks else 0
cpl_meta = (spend / meta_leads) if meta_leads else 0
cpl_ghl = (spend / ghl_meta_leads) if ghl_meta_leads else 0
cpb = (spend / ghl_booked) if ghl_booked else 0

by_acct = (daily.groupby("account_label")
           .agg(spend=("spend_aud", "sum"), impr=("impr", "sum"), clicks=("clicks", "sum"), leads=("leads", "sum"))
           .reset_index())
by_camp = (daily.groupby(["campaign_name", "account_label"])
           .agg(spend=("spend_aud", "sum"), impr=("impr", "sum"), clicks=("clicks", "sum"), leads=("leads", "sum"))
           .reset_index().sort_values("spend", ascending=False))
by_day = daily.groupby("date").agg(spend=("spend_aud", "sum"), leads=("leads", "sum")).reset_index().sort_values("date")


def fig_png(fig):
    b = io.BytesIO(); fig.savefig(b, format="png", dpi=150, bbox_inches="tight"); plt.close(fig); b.seek(0); return b


def chart_daily():
    fig, ax = plt.subplots(figsize=(8, 3))
    x = [d.strftime("%b %d") for d in by_day["date"]]
    ax.bar(x, by_day["spend"], color=BLUE)
    ax.set_title("Daily ad spend (AUD)", fontsize=12, fontweight="bold", loc="left")
    for i, v in enumerate(by_day["spend"]):
        ax.text(i, v, f"${v:,.0f}", ha="center", va="bottom", fontsize=8)
    ax.spines[["top", "right"]].set_visible(False); ax.tick_params(labelsize=8)
    return fig_png(fig)


def chart_campaigns():
    d = by_camp.copy().iloc[::-1]
    fig, ax = plt.subplots(figsize=(8, max(2.4, 0.5 * len(d))))
    colors = [PURPLE if a == "Melbourne" else GREEN for a in d["account_label"]]
    ax.barh([c[:42] for c in d["campaign_name"]], d["spend"], color=colors)
    ax.set_title("Spend by campaign (AUD)", fontsize=12, fontweight="bold", loc="left")
    for i, v in enumerate(d["spend"]):
        ax.text(v, i, f" ${v:,.0f}", va="center", fontsize=8)
    ax.spines[["top", "right"]].set_visible(False); ax.tick_params(labelsize=8)
    import matplotlib.patches as mp
    ax.legend([mp.Patch(color=PURPLE), mp.Patch(color=GREEN)], ["Melbourne", "Sydney"], fontsize=8, frameon=False)
    return fig_png(fig)


def add_table(doc, df, money_cols=(), pct_cols=(), int_cols=()):
    t = doc.add_table(rows=1, cols=len(df.columns))
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]; cell.text = str(c)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True; run.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(df.columns):
            v = row[c]
            if c in money_cols:
                txt = f"${float(v):,.0f}"
            elif c in pct_cols:
                txt = f"{float(v)*100:.2f}%"
            elif c in int_cols:
                txt = f"{int(v):,}"
            else:
                txt = str(v)
            cells[i].text = txt
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


doc = Document()
title = doc.add_paragraph()
r = title.add_run("The Migration — Meta Ads Performance")
r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
sub = doc.add_paragraph()
sub.add_run(f"{MONTH_LABEL}  (month-to-date: Jun 01 – {last_day})   ·   "
            f"Generated {datetime.now():%d %b %Y %H:%M}   ·   USD→AUD {FX:.3f}").italic = True

doc.add_heading("At a glance", 1)
kpi = pd.DataFrame([
    ("Ad Spend (AUD)", f"${spend:,.0f}"),
    ("Impressions", f"{impr:,.0f}"),
    ("Clicks", f"{clicks:,.0f}"),
    ("CTR", f"{ctr*100:.2f}%"),
    ("CPC", f"${cpc:,.2f}"),
    ("Meta-reported leads", f"{meta_leads:,}"),
    ("GHL Meta leads", f"{ghl_meta_leads:,}"),
    ("CPL (GHL Meta)", f"${cpl_ghl:,.0f}" if cpl_ghl else "—"),
    ("Booked (GHL)", f"{ghl_booked:,}"),
    ("Showed (GHL)", f"{ghl_showed:,}"),
    ("Cost per Booking", f"${cpb:,.0f}" if cpb else "—"),
], columns=["Metric", "Value"])
add_table(doc, kpi)

doc.add_heading("Daily spend", 1)
doc.add_picture(chart_daily(), width=Inches(6.3))

doc.add_heading("Spend by campaign", 1)
doc.add_picture(chart_campaigns(), width=Inches(6.3))
ct = by_camp.copy()
ct["CTR"] = (ct["clicks"] / ct["impr"]).fillna(0)
ct["CPL"] = (ct["spend"] / ct["leads"]).replace([float("inf")], 0).fillna(0)
disp = pd.DataFrame({
    "Campaign": ct["campaign_name"].str.slice(0, 46), "Acct": ct["account_label"],
    "Spend": ct["spend"], "Impr": ct["impr"], "Clicks": ct["clicks"],
    "CTR": ct["CTR"], "Leads": ct["leads"].astype(int),
    "CPL": ct["CPL"].map(lambda v: f"${v:,.0f}" if v else "—")})
add_table(doc, disp, money_cols=["Spend"], pct_cols=["CTR"], int_cols=["Impr", "Clicks"])

doc.add_heading("By account", 1)
ba = by_acct.copy()
ba["CTR"] = (ba["clicks"] / ba["impr"]).fillna(0)
bad = pd.DataFrame({"Account": ba["account_label"], "Spend": ba["spend"], "Impr": ba["impr"],
                    "Clicks": ba["clicks"], "CTR": ba["CTR"], "Meta Leads": ba["leads"].astype(int)})
add_table(doc, bad, money_cols=["Spend"], pct_cols=["CTR"], int_cols=["Impr", "Clicks"])

doc.add_heading("Funnel — ad spend to consultation (GHL)", 1)
fn = pd.DataFrame([
    ("Meta-reported leads (pixel/forms)", meta_leads),
    ("GHL Meta leads (created/revived, Meta-attributed)", ghl_meta_leads),
    ("Booked an appointment", ghl_booked),
    ("Showed", ghl_showed),
], columns=["Stage", "Count"])
add_table(doc, fn, int_cols=["Count"])

doc.add_heading("Notes", 1)
doc.add_paragraph(
    "Spend, impressions, clicks and Meta-reported leads come from the Meta Ads API (both ad accounts), "
    f"converted USD→AUD at {FX:.3f}. GHL Meta leads = contacts created or revived in the window whose "
    "latest-form source is Meta (incl. Facebook-Messenger leads identified via the Conversations channel). "
    "Booked/Showed are real GHL appointment outcomes. This is month-to-date — June data currently runs "
    f"through {last_day}.")

out = ROOT.parent / f"Meta_Ads_Performance_{MONTH_LABEL.replace(' ', '_')}.docx"
doc.save(str(out))
print(f"SAVED: {out}")
print(f"Spend AUD ${spend:,.0f} | Meta leads {meta_leads} | GHL Meta leads {ghl_meta_leads} | "
      f"booked {ghl_booked} | showed {ghl_showed} | campaigns {len(by_camp)}")
