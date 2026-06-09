"""
Build a Word (.docx) "Lead Journey — Performance Analysis" report from the
live journey DataFrames (vw_journey_leads + vw_journey_stage_cells).

Charts are rendered with matplotlib (Agg) and embedded as PNGs; every table is
data-driven from the same window the dashboard is showing, so the report is
"dynamic" — regenerate it any time and it reflects current data.

Public entry point:  build_docx(j_leads, j_cells, date_from, date_to, generated_at) -> bytes
"""
from __future__ import annotations

import io

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PIPE_RANK = {"L2C - Education": 0, "L2C - VISA": 1,
             "CLT - Onshore Admission": 2, "CLT - Admissions Sub-Applications": 3}
PIPE_COLORS = {"L2C - Education": "#3b82f6", "L2C - VISA": "#8b5cf6",
               "CLT - Onshore Admission": "#10b981", "CLT - Admissions Sub-Applications": "#f59e0b"}
SHORT = {"L2C - Education": "Education", "L2C - VISA": "VISA",
         "CLT - Onshore Admission": "Onshore", "CLT - Admissions Sub-Applications": "Sub-Apps"}


# ---------------------------------------------------------------------
# Computations (mirror the dashboard tab)
# ---------------------------------------------------------------------
def _compute(j_leads: pd.DataFrame, j_cells: pd.DataFrame) -> dict:
    tot = len(j_leads)
    edu_ord = (j_cells[j_cells["pipeline_name"] == "L2C - Education"]
               .groupby("contact_id")["stage_order"].max())
    clt_ord = (j_cells[j_cells["pipeline_name"] == "CLT - Onshore Admission"]
               .groupby("contact_id")["stage_order"].max())
    coe_ids = (set(edu_ord[edu_ord >= 11].index) | set(clt_ord[clt_ord >= 9].index)
               | set(j_leads.loc[j_leads["cur_stage"] == "COE Received", "contact_id"]))

    def pct(n):
        return f"{n / tot * 100:.0f}%" if tot else "-"

    def avg(s):
        s = s.dropna()
        return f"{s.mean():.0f}" if len(s) else "-"

    n_qual = int((edu_ord >= 1).sum())
    n_bls = int((edu_ord >= 4).sum())
    n_book = int(j_leads["days_lead_to_appt"].notna().sum())
    n_show = int(j_leads["days_lead_to_show"].notna().sum())
    n_pay = int((j_leads["total_payment"] > 0).sum())
    n_coe = len(coe_ids)

    funnel = pd.DataFrame([
        ("All Leads", tot, "100%", "-"),
        ("Reached Qualifier", n_qual, pct(n_qual), "-"),
        ("Booking Link Shared +", n_bls, pct(n_bls), "-"),
        ("Appointment Booked", n_book, pct(n_book), avg(j_leads["days_lead_to_appt"])),
        ("Consultation Showed", n_show, pct(n_show), avg(j_leads["days_lead_to_show"])),
        ("Payment Made", n_pay, pct(n_pay), avg(j_leads["days_lead_to_pay"])),
        ("COE Received", n_coe, pct(n_coe), "-"),
    ], columns=["Stage", "Leads", "% of leads", "Avg days to reach"])

    booked = j_leads[j_leads["days_lead_to_appt"].notna()]
    shown = j_leads[j_leads["days_lead_to_show"].notna()]
    paid = j_leads[j_leads["days_lead_to_pay"].notna()]
    seg_show = (shown["days_lead_to_show"] - shown["days_lead_to_appt"]).dropna()
    seg_pay = (paid["days_lead_to_pay"] - paid["days_lead_to_show"]).dropna()

    def crow(name, n, denom, days):
        return (name, int(n), f"{n / denom * 100:.0f}%" if denom else "-",
                f"{days.mean():.0f}" if len(days) else "-",
                f"{days.median():.0f}" if len(days) else "-")
    conv = pd.DataFrame([
        crow("Lead -> Booked", len(booked), tot, booked["days_lead_to_appt"].dropna()),
        crow("Booked -> Showed", len(shown), len(booked), seg_show),
        crow("Showed -> Paid", len(paid), len(shown), seg_pay),
    ], columns=["Transition", "Leads", "Conversion", "Avg days", "Median days"])

    perf = (j_cells.groupby(["pipeline_name", "stage_name", "stage_order"])
            .agg(leads=("contact_id", "nunique"),
                 avg_days=("days_in_stage", "mean"),
                 median_days=("days_in_stage", "median"),
                 max_days=("days_in_stage", "max"),
                 stuck_30=("days_in_stage", lambda s: int((s > 30).sum())))
            .reset_index())
    perf["prank"] = perf["pipeline_name"].map(PIPE_RANK).fillna(9)
    perf = perf.sort_values(["prank", "stage_order"]).reset_index(drop=True)

    return dict(tot=tot, n_qual=n_qual, n_bls=n_bls, n_book=n_book, n_show=n_show,
                n_pay=n_pay, n_coe=n_coe, funnel=funnel, conv=conv, perf=perf)


# ---------------------------------------------------------------------
# Charts -> PNG bytes
# ---------------------------------------------------------------------
def _fig_png(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_avg_days(perf: pd.DataFrame) -> io.BytesIO:
    d = perf[perf["leads"] > 0].copy()
    d["label"] = d["pipeline_name"].map(SHORT) + " | " + d["stage_name"]
    d = d.iloc[::-1]  # so first stage is at top
    colors = [PIPE_COLORS.get(p, "#9ca3af") for p in d["pipeline_name"]]
    fig, ax = plt.subplots(figsize=(8.2, max(3.5, 0.30 * len(d))))
    ax.barh(d["label"], d["avg_days"], color=colors)
    ax.set_xlabel("Avg days in stage (today - last update)")
    ax.set_title("Average days in stage", fontsize=12, fontweight="bold", loc="left")
    for i, v in enumerate(d["avg_days"]):
        ax.text(v + 0.5, i, f"{v:.0f}", va="center", fontsize=7.5, color="#333")
    handles = [plt.Rectangle((0, 0), 1, 1, color=c) for c in PIPE_COLORS.values()]
    ax.legend(handles, [SHORT[p] for p in PIPE_COLORS], loc="lower right", fontsize=8, frameon=False)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(labelsize=8)
    return _fig_png(fig)


def _chart_funnel(funnel: pd.DataFrame) -> io.BytesIO:
    d = funnel.copy()
    fig, ax = plt.subplots(figsize=(8.2, 3.6))
    bars = ax.bar(d["Stage"], d["Leads"], color="#3b82f6")
    ax.set_title("Journey funnel — leads reaching each milestone",
                 fontsize=12, fontweight="bold", loc="left")
    ax.set_ylabel("Leads")
    for b, lab in zip(bars, d["% of leads"]):
        ax.text(b.get_x() + b.get_width() / 2, b.get_height(),
                f"{int(b.get_height())}\n{lab}", ha="center", va="bottom", fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)
    ax.tick_params(axis="x", labelrotation=30, labelsize=8)
    for t in ax.get_xticklabels():
        t.set_ha("right")
    return _fig_png(fig)


# ---------------------------------------------------------------------
# Word helpers
# ---------------------------------------------------------------------
def _add_table(doc, df: pd.DataFrame):
    t = doc.add_table(rows=1, cols=len(df.columns))
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(df.columns):
            v = row[c]
            cells[i].text = "" if (v is None or (isinstance(v, float) and pd.isna(v))) else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t


def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


# ---------------------------------------------------------------------
# Public builder
# ---------------------------------------------------------------------
def build_docx(j_leads: pd.DataFrame, j_cells: pd.DataFrame,
               date_from, date_to, generated_at: str = "") -> bytes:
    m = _compute(j_leads, j_cells)
    doc = Document()

    # Title block
    title = doc.add_paragraph()
    run = title.add_run("The Migration — Lead Journey Performance Analysis")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.add_run(f"Window: {date_from:%d %b %Y} – {date_to:%d %b %Y}"
                + (f"   ·   Generated: {generated_at}" if generated_at else "")).italic = True
    doc.add_paragraph(
        "Cohort = every contact created or revived (latest form/survey submission) in the window. "
        "Pipelines tracked: L2C - Education, L2C - VISA, CLT - Onshore Admission, "
        "CLT - Admissions Sub-Applications.")

    # KPI line
    _h(doc, "At a glance", 1)
    kpi = pd.DataFrame([
        ("Total Leads", f"{m['tot']:,}"),
        ("Reached Qualifier", f"{m['n_qual']:,}"),
        ("Appointment Booked", f"{m['n_book']:,}"),
        ("Consultation Showed", f"{m['n_show']:,}"),
        ("Payment Made", f"{m['n_pay']:,}"),
        ("COE Received", f"{m['n_coe']:,}"),
    ], columns=["Metric", "Value"])
    _add_table(doc, kpi)

    # Chart 1 + by-stage table
    _h(doc, "Where leads sit longest (by stage)", 1)
    doc.add_picture(_chart_avg_days(m["perf"]), width=Inches(6.4))
    perf_disp = pd.DataFrame({
        "Pipeline": m["perf"]["pipeline_name"],
        "Stage": m["perf"]["stage_name"],
        "Leads here": m["perf"]["leads"].astype(int),
        "Avg days": m["perf"]["avg_days"].round(0).astype("Int64"),
        "Median": m["perf"]["median_days"].round(0).astype("Int64"),
        "Longest": m["perf"]["max_days"].round(0).astype("Int64"),
        "Stuck >30d": m["perf"]["stuck_30"].astype(int),
    })
    _add_table(doc, perf_disp)

    # Chart 2 + funnel table
    _h(doc, "Journey funnel — how far leads get", 1)
    doc.add_picture(_chart_funnel(m["funnel"]), width=Inches(6.4))
    _add_table(doc, m["funnel"])

    # Conversion table
    _h(doc, "Milestone conversion & speed", 1)
    _add_table(doc, m["conv"])

    # Methodology note
    _h(doc, "Notes & methodology", 1)
    doc.add_paragraph(
        "Real timestamps (exact): appointment booked, consultation shown, and first payment — "
        "these drive the funnel 'avg days to reach' and the conversion-speed table.", style=None)
    doc.add_paragraph(
        "Stage dwell ('days in stage', 'stuck >30d') = today minus the opportunity's last update — "
        "the only per-stage signal available, because GHL's sync stores no stage-change history "
        "(days_in_pipeline = 0 for every opportunity). Pure pipeline stages (Qualifier, Pre Sales, "
        "Booking Link Shared, Initial Requested/Received) therefore can't be timed individually.")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
